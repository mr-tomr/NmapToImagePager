import os
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_caption_style(paragraph):
    """
    Set the style of the caption to Calibri, 9 points, and centered.
    """
    run = paragraph.runs[0]
    run.font.name = 'Calibri'
    run.font.size = Pt(9)

    # Ensure the font is set to Calibri
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_images_to_word(folder, output_filename):
    """
    Add all images from the specified folder to a Word document.
    """
    if not os.path.isdir(folder):
        print(f"Error: The folder '{folder}' does not exist.")
        return

    # Create a new Word document
    document = Document()

    # List all image files in the folder
    image_files = [f for f in os.listdir(folder) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif'))]
    image_files.sort()  # Ensure the images are sorted by filename

    # Add each image to the Word document
    for image_file in image_files:
        image_path = os.path.join(folder, image_file)
        # Add image and center it
        image_paragraph = document.add_paragraph()
        run = image_paragraph.add_run()
        run.add_picture(image_path, width=Inches(6))  # Adjust the width as needed
        image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add caption and set its style
        caption_paragraph = document.add_paragraph(image_file)
        set_caption_style(caption_paragraph)
    
    # Save the Word document
    document.save(output_filename)
    print(f"Images saved to {output_filename}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python script.py <image_folder> [output_filename]")
        sys.exit(1)
    
    folder = sys.argv[1]
    output_filename = sys.argv[2] if len(sys.argv) > 2 else f"{folder}.docx"
    add_images_to_word(folder, output_filename)
